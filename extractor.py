"""
Extrator de questões de provas .docx — Colégio Fleming

Recebe um .docx de professor (formatação livre) e devolve um JSON estruturado
com cada questão separada: enunciado com runs de formatação, alternativas,
referências a imagens e equações OMML.

Uso:
    python extractor.py arquivo1.docx [arquivo2.docx ...] [-o saida.json]
"""

import argparse
import base64
import copy
import json
import os
import re
import sys
from dataclasses import dataclass, field, asdict
from typing import Optional
from lxml import etree
import docx
from docx.oxml.ns import qn

NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
}

RE_QUESTION_NUM = re.compile(r"^(\d{1,3})\s*\)\s*")
RE_ALTERNATIVE = re.compile(r"^([a-eA-E])\s*[\)\-]\s*")


# ---------------------------------------------------------------------------
# Data classes for structured output
# ---------------------------------------------------------------------------

@dataclass
class RunFragment:
    type: str  # "text", "image", "equation", "line_break"
    content: str = ""
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    font_name: Optional[str] = None
    font_size_pt: Optional[float] = None
    color: Optional[str] = None
    image_filename: Optional[str] = None
    image_width_px: Optional[int] = None
    image_height_px: Optional[int] = None
    image_crop: Optional[dict] = None  # {l, t, r, b} in thousandths of percent
    omml_xml: Optional[str] = None


@dataclass
class Alternative:
    letter: str
    fragments: list = field(default_factory=list)


@dataclass
class Question:
    number: int
    original_number: int
    source_file: str = ""
    source_tag: str = ""  # e.g. "(UFSM)"
    weight: Optional[str] = None  # e.g. "1,0 ponto"
    statement: list = field(default_factory=list)  # list of RunFragment dicts
    alternatives: list = field(default_factory=list)  # list of Alternative dicts


# ---------------------------------------------------------------------------
# Image extraction
# ---------------------------------------------------------------------------

def _extract_images_from_doc(doc) -> dict:
    """Return {rId: (filename, bytes)} for all images in the document."""
    images = {}
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            blob = rel.target_part.blob
            ext = os.path.splitext(rel.target_ref)[1] or ".png"
            filename = f"{rel_id}{ext}"
            images[rel_id] = (filename, blob)
    return images


def _find_image_in_element(el, images_map) -> Optional[RunFragment]:
    """Look for a drawing/inline image inside an XML element."""
    for drawing in el.findall(".//" + qn("w:drawing")):
        blip = drawing.find(".//" + qn("a:blip"))
        if blip is None:
            continue
        r_embed = blip.get(qn("r:embed"))
        if r_embed and r_embed in images_map:
            fname, blob = images_map[r_embed]
            extent = drawing.find(".//" + qn("wp:extent"))
            w_px = h_px = None
            if extent is not None:
                cx = int(extent.get("cx", 0))
                cy = int(extent.get("cy", 0))
                w_px = int(cx / 9525) if cx else None
                h_px = int(cy / 9525) if cy else None
            crop = None
            src_rect = drawing.find(".//" + qn("a:srcRect"))
            if src_rect is not None:
                cl = int(src_rect.get("l", "0"))
                ct = int(src_rect.get("t", "0"))
                cr = int(src_rect.get("r", "0"))
                cb = int(src_rect.get("b", "0"))
                if cl or ct or cr or cb:
                    crop = {"l": cl, "t": ct, "r": cr, "b": cb}
            return RunFragment(
                type="image",
                image_filename=fname,
                image_width_px=w_px,
                image_height_px=h_px,
                image_crop=crop,
            )
    return None


# ---------------------------------------------------------------------------
# Paragraph → list of RunFragment
# ---------------------------------------------------------------------------

def _extract_paragraph_fragments(para, images_map) -> list:
    """Walk the paragraph XML children in order, producing RunFragments."""
    fragments = []
    p_el = para._element

    for child in p_el:
        tag = etree.QName(child).localname

        if tag == "r":  # regular run
            img_frag = _find_image_in_element(child, images_map)
            if img_frag:
                fragments.append(img_frag)
                continue

            # Check for line break
            if child.find(qn("w:br")) is not None:
                fragments.append(RunFragment(type="line_break"))

            t_el = child.find(qn("w:t"))
            if t_el is not None and t_el.text:
                rpr = child.find(qn("w:rPr"))
                frag = RunFragment(type="text", content=t_el.text)
                if rpr is not None:
                    b = rpr.find(qn("w:b"))
                    if b is not None and b.get(qn("w:val"), "true") != "false":
                        frag.bold = True
                    i = rpr.find(qn("w:i"))
                    if i is not None and i.get(qn("w:val"), "true") != "false":
                        frag.italic = True
                    u = rpr.find(qn("w:u"))
                    if u is not None:
                        val = u.get(qn("w:val"), "")
                        if val and val != "none":
                            frag.underline = True
                    rfonts = rpr.find(qn("w:rFonts"))
                    if rfonts is not None:
                        frag.font_name = (
                            rfonts.get(qn("w:ascii"))
                            or rfonts.get(qn("w:hAnsi"))
                        )
                    sz = rpr.find(qn("w:sz"))
                    if sz is not None:
                        val = sz.get(qn("w:val"))
                        if val:
                            frag.font_size_pt = int(val) / 2
                    color = rpr.find(qn("w:color"))
                    if color is not None:
                        frag.color = color.get(qn("w:val"))
                fragments.append(frag)

        elif tag == "oMath" or tag == "oMathPara":
            omml_str = etree.tostring(child, encoding="unicode")
            text = _omml_to_text(child)
            fragments.append(RunFragment(
                type="equation",
                content=text,
                omml_xml=omml_str,
            ))

        elif tag == "hyperlink":
            for r in child.findall(qn("w:r")):
                t_el = r.find(qn("w:t"))
                if t_el is not None and t_el.text:
                    fragments.append(RunFragment(type="text", content=t_el.text))

    return fragments


def _omml_to_text(omml_el) -> str:
    """Best-effort plain-text rendering of OMML for preview purposes."""
    texts = []
    for t in omml_el.findall(".//" + qn("m:t")):
        if t.text:
            texts.append(t.text)
    return "".join(texts)


# ---------------------------------------------------------------------------
# Question detection and grouping
# ---------------------------------------------------------------------------

def _get_plain_text(fragments: list) -> str:
    return "".join(f.get("content", "") if isinstance(f, dict) else f.content
                   for f in fragments
                   if (f.get("type") if isinstance(f, dict) else f.type) in ("text", "equation"))


def _detect_question_start(text: str):
    """If text starts with 'N) ...', return (N, rest_of_text)."""
    m = RE_QUESTION_NUM.match(text.strip())
    if m:
        return int(m.group(1))
    return None


def _detect_alternative(text: str):
    """If text starts with 'a) ...' or 'A - ...', return letter."""
    m = RE_ALTERNATIVE.match(text.strip())
    if m:
        return m.group(1).lower()
    return None


def _detect_list_paragraph_numbering(para) -> Optional[str]:
    """Detect Word auto-numbered list items (a), b), c)...) from XML numbering."""
    p_el = para._element
    num_pr = p_el.find(qn("w:pPr") + "/" + qn("w:numPr"))
    if num_pr is None:
        return None
    style = para.style.name if para.style else ""
    if "List" not in style:
        return None
    return True  # signal that this is a numbered list item


def _detect_weight(text: str) -> Optional[str]:
    """Detect '(vale X ponto)' annotations."""
    m = re.search(r"\(vale\s+([\d,\.]+)\s*pontos?\)", text, re.IGNORECASE)
    if m:
        return m.group(1) + " ponto"
    return None


def _detect_source_tag(text: str) -> str:
    """Detect '(UFSM)' or similar attribution at start of question."""
    m = re.search(r"\(([A-Z]{2,}(?:\s*[-/]\s*[A-Z]+)*)\)", text[:100])
    if m:
        return f"({m.group(1)})"
    return ""


def _is_header_or_metadata(text: str) -> bool:
    """Detect header rows, empty lines, formulário, gabarito labels."""
    t = text.strip().upper()
    if not t:
        return True
    if RE_QUESTION_NUM.match(t):
        return False
    skip_patterns = [
        "COLÉGIO FLEMING", "ALUNO(A):", "PROFESSOR (A):", "PROFESSOR:",
        "COMPONENTE CURRICULAR", "SÉRIE:", "TURMA:", "DATA:", "PESO:",
        "NOTA:", "GABARITO:", "GABARITO",
        "FORMULÁRIO", "QUESTÕES PROVA",
    ]
    if any(p in t for p in skip_patterns):
        if len(t) < 200:
            return True
    return False


def _is_formula_line(text: str) -> bool:
    """Detect formulário lines (physics/math formulas before questions)."""
    t = text.strip()
    if RE_QUESTION_NUM.match(t):
        return False
    if RE_ALTERNATIVE.match(t):
        return False
    if len(t) > 120:
        return False
    t_no_space = t.replace(" ", "")
    formula_indicators = ["W=F", "P=m.g", "FR=m", "F.d.cos",
                          "S=S0", "V=V0", "V²=", "E_C=", "Ec=",
                          "E_PG=", "Epg=", "E_PE=", "Epe="]
    hits = sum(1 for ind in formula_indicators if ind in t_no_space)
    return hits >= 1 and len(t) < 80


def _serialize_fragment(frag: RunFragment) -> dict:
    """Convert RunFragment to dict, dropping None values for compact JSON."""
    d = {}
    d["type"] = frag.type
    if frag.content:
        d["content"] = frag.content
    if frag.bold:
        d["bold"] = True
    if frag.italic:
        d["italic"] = True
    if frag.underline:
        d["underline"] = True
    if frag.font_name:
        d["font"] = frag.font_name
    if frag.font_size_pt:
        d["size_pt"] = frag.font_size_pt
    if frag.color and frag.color != "000000":
        d["color"] = frag.color
    if frag.image_filename:
        d["image"] = frag.image_filename
    if frag.image_width_px:
        d["width_px"] = frag.image_width_px
    if frag.image_height_px:
        d["height_px"] = frag.image_height_px
    if frag.image_crop:
        d["crop"] = frag.image_crop
    if frag.omml_xml:
        d["omml"] = frag.omml_xml
    return d


# ---------------------------------------------------------------------------
# Main extraction logic
# ---------------------------------------------------------------------------

def extract_questions(docx_path: str) -> dict:
    """
    Extract all questions from a .docx file.

    Returns a dict with:
      - source_file: filename
      - metadata: detected header info
      - formulario: formula fragments (if present)
      - questions: list of Question dicts
      - images: {filename: base64_data}
    """
    doc = docx.Document(docx_path)
    images_map = _extract_images_from_doc(doc)
    source_file = os.path.basename(docx_path)

    # Extract all paragraphs with their fragments
    para_data = []
    for para in doc.paragraphs:
        frags = _extract_paragraph_fragments(para, images_map)
        plain = _get_plain_text([_serialize_fragment(f) for f in frags])
        para_data.append({
            "fragments": frags,
            "plain_text": plain,
            "style": para.style.name if para.style else "",
            "alignment": str(para.alignment) if para.alignment else "default",
        })

    # Classify paragraphs
    questions = []
    current_question = None
    current_alternative = None
    formulario_frags = []
    in_formulario = False
    list_alt_counter = 0  # for auto-numbered list alternatives
    metadata = {"professors": [], "subject": "", "series": "", "type": ""}

    # Try to extract metadata from tables (header table)
    for table in doc.tables:
        for row in table.rows:
            unique_cells = []
            seen = set()
            for c in row.cells:
                if c.text.strip() not in seen:
                    unique_cells.append(c.text.strip())
                    seen.add(c.text.strip())
            row_text = " ".join(unique_cells)
            if "Professor" in row_text:
                prof_match = re.search(r"Professor\s*\(a\)\s*:\s*(.+)", row_text)
                if prof_match:
                    metadata["professors"] = [
                        p.strip() for p in prof_match.group(1).split("/")
                    ]
            if "Componente" in row_text:
                comp_match = re.search(r"Componente\s+Curricular\s*:\s*(\w+)", row_text)
                if comp_match:
                    metadata["subject"] = comp_match.group(1)
            if "Série" in row_text:
                serie_match = re.search(r"Série\s*:\s*(\S+)", row_text)
                if serie_match:
                    metadata["series"] = serie_match.group(1)
            if "Tipo" in row_text:
                tipo_match = re.search(r"Tipo\s*:\s*(.+?)(?:Série|$)", row_text)
                if tipo_match:
                    metadata["type"] = tipo_match.group(1).strip()

    for idx, pd in enumerate(para_data):
        plain = pd["plain_text"]
        frags = pd["fragments"]
        para_obj = doc.paragraphs[idx]

        # Skip empty / header / metadata — but keep paragraphs with images
        has_images = any(
            (f.type if hasattr(f, 'type') else f.get('type')) == 'image'
            for f in frags
        )
        if _is_header_or_metadata(plain) and not has_images:
            list_alt_counter = 0
            continue

        # Detect formulário section
        if plain.strip().upper() == "FORMULÁRIO" or _is_formula_line(plain):
            in_formulario = True
            formulario_frags.extend([_serialize_fragment(f) for f in frags])
            list_alt_counter = 0
            continue

        # Check for question start
        q_num = _detect_question_start(plain)
        if q_num is not None:
            in_formulario = False
            list_alt_counter = 0
            # Save previous question
            if current_question:
                if current_alternative:
                    current_question.alternatives.append(
                        asdict(current_alternative)
                    )
                    current_alternative = None
                questions.append(current_question)

            current_question = Question(
                number=q_num,
                original_number=q_num,
                source_file=source_file,
                source_tag=_detect_source_tag(plain),
                weight=_detect_weight(plain),
            )
            current_question.statement = [_serialize_fragment(f) for f in frags]
            current_alternative = None
            continue

        if in_formulario:
            formulario_frags.extend([_serialize_fragment(f) for f in frags])
            continue

        if current_question is None:
            continue

        # Check for alternative
        alt_letter = _detect_alternative(plain)
        if alt_letter:
            list_alt_counter = 0
            if current_alternative:
                current_question.alternatives.append(asdict(current_alternative))
            alt_frags = _strip_letter_prefix(frags, alt_letter)
            current_alternative = Alternative(
                letter=alt_letter,
                fragments=[_serialize_fragment(f) for f in alt_frags],
            )

            # Check if multiple alternatives are on the same line
            if _count_alternatives_in_line(plain) > 1:
                alts = _split_inline_alternatives(plain, frags)
                if alts:
                    current_alternative = None
                    for a in alts:
                        current_question.alternatives.append(asdict(a))
            continue

        # Check for auto-numbered list paragraph (Word adds a/b/c automatically)
        is_list = _detect_list_paragraph_numbering(para_obj)
        if is_list and current_question and not alt_letter:
            letters = "abcde"
            if list_alt_counter < len(letters):
                letter = letters[list_alt_counter]
                list_alt_counter += 1
                if current_alternative:
                    current_question.alternatives.append(asdict(current_alternative))
                current_alternative = Alternative(
                    letter=letter,
                    fragments=[_serialize_fragment(f) for f in frags],
                )
                continue

        list_alt_counter = 0

        # Continuation of current context (question statement or alternative)
        serialized = [_serialize_fragment(f) for f in frags]
        if current_alternative:
            current_alternative.fragments.extend(serialized)
        else:
            current_question.statement.extend(serialized)

    # Save last question
    if current_question:
        if current_alternative:
            current_question.alternatives.append(asdict(current_alternative))
        questions.append(current_question)

    # Build image map (base64)
    images_b64 = {}
    for rid, (fname, blob) in images_map.items():
        images_b64[fname] = base64.b64encode(blob).decode("ascii")

    return {
        "source_file": source_file,
        "metadata": metadata,
        "formulario": formulario_frags,
        "questions": [asdict(q) for q in questions],
        "images": images_b64,
        "stats": {
            "total_questions": len(questions),
            "total_images": len(images_b64),
            "total_equations": sum(
                1 for q in questions
                for f in q.statement
                if isinstance(f, dict) and f.get("type") == "equation"
            ),
        },
    }


def _strip_letter_prefix(frags: list, letter: str) -> list:
    """Remove the 'a) ' or 'A - ' prefix from the first text fragment(s).

    Handles both single-run ('a) text') and split-run ('a' + ') text') patterns.
    """
    result = list(frags)
    for i, f in enumerate(result):
        if f.type != "text" or not f.content:
            continue

        # Case 1: "a) text" in one run
        m = re.match(
            rf"^\s*{re.escape(letter)}\s*[\)\-]\s*",
            f.content, re.IGNORECASE
        )
        if m:
            new_f = copy.copy(f)
            new_f.content = f.content[m.end():]
            result[i] = new_f
            break

        # Case 2: first run is just the letter "a"
        if f.content.strip().lower() == letter.lower():
            result.pop(i)
            # Remove ") " from next text fragment
            for j in range(i, len(result)):
                if result[j].type == "text" and result[j].content:
                    m2 = re.match(r"^\s*[\)\-]\s*", result[j].content)
                    if m2:
                        new_f = copy.copy(result[j])
                        new_f.content = result[j].content[m2.end():]
                        result[j] = new_f
                    break
            break

    return result


def _count_alternatives_in_line(text: str) -> int:
    """Count how many alternatives appear on the same line (only at word boundary)."""
    return len(re.findall(r"(?:^|(?<=\s))[a-eA-E]\s*[\)\-]", text))


def _split_inline_alternatives(plain_text: str, frags: list) -> list:
    """
    Split a line like 'a)0,5 W.  b)5,0 W.  c)50 W.' into separate alternatives.
    Uses the plain text to find boundaries, then maps back to fragments.
    """
    full_text = plain_text.strip()
    matches = list(re.finditer(r"(?:^|(?<=\s))([a-eA-E])\s*[\)\-]\s*", full_text))
    if len(matches) < 2:
        return []

    alternatives = []
    for i, m in enumerate(matches):
        letter = m.group(1).lower()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(full_text)
        content = full_text[start:end].strip()
        alt = Alternative(
            letter=letter,
            fragments=[{"type": "text", "content": content}],
        )
        alternatives.append(alt)

    return alternatives


# ---------------------------------------------------------------------------
# Merge + renumber
# ---------------------------------------------------------------------------

def merge_extractions(extractions: list, renumber: bool = True) -> dict:
    """
    Merge multiple extraction results into a single question set.
    Renumbers questions sequentially if renumber=True.
    """
    all_questions = []
    all_images = {}
    all_formulario = []
    sources = []

    for ext in extractions:
        sources.append(ext["source_file"])
        offset = len(all_questions)

        # Prefix image filenames to avoid collisions
        prefix = f"src{len(sources)}_"
        img_remap = {}
        for fname, b64 in ext["images"].items():
            new_fname = prefix + fname
            all_images[new_fname] = b64
            img_remap[fname] = new_fname

        for q in ext["questions"]:
            q_copy = copy.deepcopy(q)

            # Remap image references
            _remap_images_in_fragments(q_copy["statement"], img_remap)
            for alt in q_copy["alternatives"]:
                _remap_images_in_fragments(alt["fragments"], img_remap)

            if renumber:
                old_num = q_copy["number"]
                new_num = offset + (old_num - min(qq["number"] for qq in ext["questions"])) + 1
                q_copy["original_number"] = old_num
                q_copy["number"] = len(all_questions) + 1
                _renumber_in_fragments(q_copy["statement"], old_num, q_copy["number"])

            all_questions.append(q_copy)

        if ext["formulario"] and not all_formulario:
            all_formulario = ext["formulario"]

    return {
        "source_files": sources,
        "formulario": all_formulario,
        "questions": all_questions,
        "images": all_images,
        "stats": {
            "total_questions": len(all_questions),
            "total_images": len(all_images),
            "source_count": len(sources),
        },
    }


def _remap_images_in_fragments(fragments: list, remap: dict):
    for f in fragments:
        if f.get("image") in remap:
            f["image"] = remap[f["image"]]


def _renumber_in_fragments(fragments: list, old_num: int, new_num: int):
    """Update the question number in the first text fragment(s).

    Handles both single-run ('3) text...') and split-run ('3' + ')' + ' text')
    patterns from Word docs.
    """
    old_str = str(old_num)
    new_str = str(new_num)

    for i, f in enumerate(fragments):
        if f.get("type") != "text" or not f.get("content"):
            continue
        content = f["content"]
        stripped = content.lstrip()

        # Case 1: "3) text..." in one run
        m = re.match(rf"^(\s*){old_num}\s*\)", content)
        if m:
            f["content"] = re.sub(
                rf"^(\s*){old_num}(\s*\))", rf"\g<1>{new_num}\2", content, count=1
            )
            return

        # Case 2: first run is just the number "3", next run starts with ")"
        if stripped == old_str:
            f["content"] = content.replace(old_str, new_str, 1)
            return

        # Case 3: run starts with "N) " at the very beginning
        if stripped.startswith(old_str + ")"):
            f["content"] = content.replace(old_str + ")", new_str + ")", 1)
            return


# ---------------------------------------------------------------------------
# Human-readable summary
# ---------------------------------------------------------------------------

def print_summary(result: dict):
    """Print a human-readable summary of extracted questions."""
    src = result.get("source_file") or ", ".join(result.get("source_files", []))
    print(f"\n{'='*60}")
    print(f"Fonte: {src}")
    print(f"{'='*60}")

    meta = result.get("metadata", {})
    if meta.get("subject"):
        print(f"Disciplina: {meta['subject']}")
    if meta.get("professors"):
        print(f"Professores: {', '.join(meta['professors'])}")
    if meta.get("series"):
        print(f"Série: {meta['series']}")

    if result.get("formulario"):
        n_formulas = sum(1 for f in result["formulario"] if f.get("type") == "equation")
        print(f"\nFormulário: {n_formulas} equações detectadas")

    print(f"\nTotal de questões: {result['stats']['total_questions']}")
    print(f"Total de imagens: {result['stats']['total_images']}")

    for q in result["questions"]:
        plain_stmt = "".join(
            f.get("content", "") for f in q["statement"]
            if f.get("type") in ("text", "equation")
        )[:100]
        n_alts = len(q["alternatives"])
        has_img = any(
            f.get("type") == "image"
            for f in q["statement"]
        )
        has_eq = any(
            f.get("type") == "equation"
            for f in q["statement"]
        )
        tags = []
        if has_img:
            tags.append("IMG")
        if has_eq:
            tags.append("EQ")
        if q.get("source_tag"):
            tags.append(q["source_tag"])
        if q.get("weight"):
            tags.append(q["weight"])

        tag_str = f" [{', '.join(tags)}]" if tags else ""
        orig = ""
        if q["original_number"] != q["number"]:
            orig = f" (era Q{q['original_number']} de {q['source_file']})"

        print(f"\n  Q{q['number']}{tag_str}{orig}")
        print(f"    {plain_stmt}...")
        print(f"    → {n_alts} alternativas: {', '.join(a['letter'] for a in q['alternatives'])}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Extrai questões de provas .docx do Colégio Fleming"
    )
    parser.add_argument("files", nargs="+", help="Arquivo(s) .docx de entrada")
    parser.add_argument("-o", "--output", help="Arquivo JSON de saída")
    parser.add_argument("--merge", action="store_true",
                        help="Junta questões de múltiplos arquivos e renumera")
    parser.add_argument("--summary", action="store_true", default=True,
                        help="Mostra resumo no terminal")
    parser.add_argument("--no-images", action="store_true",
                        help="Não inclui base64 das imagens no JSON (mais leve)")

    args = parser.parse_args()

    extractions = []
    for fpath in args.files:
        print(f"Processando: {fpath}")
        result = extract_questions(fpath)
        extractions.append(result)

        if args.summary and not args.merge:
            print_summary(result)

    if args.merge and len(extractions) > 1:
        merged = merge_extractions(extractions, renumber=True)
        if args.summary:
            print_summary(merged)
        output_data = merged
    elif len(extractions) == 1:
        output_data = extractions[0]
    else:
        output_data = {"extractions": extractions}

    if args.no_images:
        if "images" in output_data:
            output_data["images"] = {k: "<base64 omitido>" for k in output_data["images"]}

    if args.output:
        with open(args.output, "w", encoding="utf-8") as f:
            json.dump(output_data, f, ensure_ascii=False, indent=2)
        print(f"\nJSON salvo em: {args.output}")


if __name__ == "__main__":
    main()
