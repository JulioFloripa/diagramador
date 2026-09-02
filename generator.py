"""
Gerador de provas .docx padronizadas — Colégio Fleming

Recebe o JSON produzido pelo extractor.py e gera um .docx formatado
no padrão institucional do Fleming:
  1. Cabeçalho (tabela com logo, dados da prova, campos do aluno)
  2. Formulário (se houver)
  3. Questões numeradas com alternativas
  4. Grade de gabarito ao final

Uso:
    python generator.py input.json -o prova.docx [--logo logo.png]
    python generator.py input.json -o prova.docx --gabarito respostas.docx
"""

import argparse
import base64
import io
import json
import os
import re
import sys
import tempfile
from typing import Optional

from PIL import Image as PILImage

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree


# ---------------------------------------------------------------------------
# Constants — Fleming institutional style
# ---------------------------------------------------------------------------

FONT_BODY = "Arial"
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_HEADER = Pt(10)
FONT_SIZE_SMALL = Pt(9)
FONT_SIZE_QUESTION_NUM = Pt(11)

PAGE_MARGIN_TOP = Cm(1.5)
PAGE_MARGIN_BOTTOM = Cm(1.5)
PAGE_MARGIN_LEFT = Cm(2.0)
PAGE_MARGIN_RIGHT = Cm(2.0)

HEADER_BG_COLOR = "1F4E79"
HEADER_TEXT_COLOR = "FFFFFF"

ALT_LETTERS = ["a", "b", "c", "d", "e"]


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, color: str):
    """Set background color of a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    cell._element.get_or_add_tcPr().append(shading)


def _set_cell_border(cell, **kwargs):
    """Set borders on a cell. kwargs: top, bottom, left, right with value (sz, color)."""
    tc_pr = cell._element.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tc_pr.append(tc_borders)
    for edge, props in kwargs.items():
        sz = props.get("sz", "4")
        color = props.get("color", "000000")
        border_el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{sz}" '
            f'w:space="0" w:color="{color}"/>'
        )
        existing = tc_borders.find(qn(f"w:{edge}"))
        if existing is not None:
            tc_borders.remove(existing)
        tc_borders.append(border_el)


def _set_cell_vertical_alignment(cell, align="center"):
    tc_pr = cell._element.get_or_add_tcPr()
    v_align = parse_xml(
        f'<w:vAlign {nsdecls("w")} w:val="{align}"/>'
    )
    existing = tc_pr.find(qn("w:vAlign"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_pr.append(v_align)


def _set_cell_width(cell, width_cm: float):
    tc_pr = cell._element.get_or_add_tcPr()
    tc_w = parse_xml(
        f'<w:tcW {nsdecls("w")} w:w="{int(width_cm * 567)}" w:type="dxa"/>'
    )
    existing = tc_pr.find(qn("w:tcW"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_pr.append(tc_w)


def _merge_cells_in_row(table, row_idx, start_col, end_col):
    """Merge cells in a row from start_col to end_col (inclusive)."""
    row = table.rows[row_idx]
    cell_a = row.cells[start_col]
    cell_b = row.cells[end_col]
    cell_a.merge(cell_b)


def _set_paragraph_spacing(para, before=0, after=0, line=None):
    """Set paragraph spacing in points."""
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = Pt(line)


def _add_run_styled(paragraph, text, bold=False, italic=False, underline=False,
                    font_name=FONT_BODY, font_size=FONT_SIZE_BODY,
                    color=None):
    """Add a styled run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if underline:
        run.underline = True
    if color:
        if isinstance(color, str):
            run.font.color.rgb = RGBColor.from_string(color)
        else:
            run.font.color.rgb = color
    return run


def _add_omml_to_paragraph(paragraph, omml_xml: str):
    """Insert an OMML equation into a paragraph."""
    try:
        omml_el = etree.fromstring(omml_xml.encode("utf-8"))
        paragraph._element.append(omml_el)
    except Exception:
        paragraph.add_run("[equação]")


def _apply_crop(image_data: bytes, crop: dict) -> bytes:
    """Apply Word-style crop to image. Crop values are in thousandths of percent."""
    img = PILImage.open(io.BytesIO(image_data))
    w, h = img.size

    left = int(w * crop.get("l", 0) / 100000)
    top = int(h * crop.get("t", 0) / 100000)
    right = w - int(w * crop.get("r", 0) / 100000)
    bottom = h - int(h * crop.get("b", 0) / 100000)

    if left >= right or top >= bottom:
        return image_data

    cropped = img.crop((left, top, right, bottom))
    buf = io.BytesIO()
    fmt = img.format or "PNG"
    cropped.save(buf, format=fmt)
    return buf.getvalue()


def _add_image_to_paragraph(paragraph, image_data: bytes, width_px=None,
                            height_px=None, crop=None, max_width_cm=14.0):
    """Add an image to a paragraph from raw bytes, applying crop if present."""
    if crop:
        image_data = _apply_crop(image_data, crop)

    stream = io.BytesIO(image_data)
    if width_px and width_px > 0:
        width_cm = width_px * 2.54 / 96
        if width_cm > max_width_cm:
            width_cm = max_width_cm
        if width_cm < 3.0:
            width_cm = min(8.0, max_width_cm)
    else:
        width_cm = min(10.0, max_width_cm)
    run = paragraph.add_run()
    run.add_picture(stream, width=Cm(width_cm))


# ---------------------------------------------------------------------------
# Header table (institutional layout)
# ---------------------------------------------------------------------------

def _build_header_table(doc, metadata: dict, logo_path: Optional[str] = None):
    """
    Build the Fleming header table:
    ┌──────────┬─────────────────────────────────────────────┐
    │  LOGO    │  COLÉGIO FLEMING                            │
    │          │  Componente: ___   Série: ___   Tipo: ___   │
    ├──────────┴─────────────────────────────────────────────┤
    │  Aluno(a): _______________  Turma: ___  Data: ___     │
    │  Professor(a): ___________  Peso: ___   Nota: ___     │
    └────────────────────────────────────────────────────────┘
    """
    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    _set_table_borders(table)

    # Row 0: Logo + School name (merged across cols 1-3)
    _merge_cells_in_row(table, 0, 1, 3)
    logo_cell = table.rows[0].cells[0]
    name_cell = table.rows[0].cells[1]

    if logo_path and os.path.exists(logo_path):
        p = logo_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(2.5))
    else:
        p = logo_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run_styled(p, "FLEMING", bold=True, font_size=Pt(14),
                       color=HEADER_BG_COLOR)

    _set_cell_shading(logo_cell, HEADER_BG_COLOR)
    _set_cell_vertical_alignment(logo_cell)

    p = name_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run_styled(p, "COLÉGIO FLEMING", bold=True, font_size=Pt(14),
                   color=HEADER_TEXT_COLOR)
    _set_cell_shading(name_cell, HEADER_BG_COLOR)
    _set_cell_vertical_alignment(name_cell)

    # Row 1: Subject info
    _merge_cells_in_row(table, 1, 0, 3)
    info_cell = table.rows[1].cells[0]
    p = info_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subject = metadata.get("subject", "_______________")
    series = metadata.get("series", "____")
    tipo = metadata.get("type", "A")

    _add_run_styled(p, "Componente Curricular: ", font_size=FONT_SIZE_HEADER)
    _add_run_styled(p, subject, bold=True, font_size=FONT_SIZE_HEADER)
    _add_run_styled(p, "     Série: ", font_size=FONT_SIZE_HEADER)
    _add_run_styled(p, series, bold=True, font_size=FONT_SIZE_HEADER)
    _add_run_styled(p, "     Tipo: ", font_size=FONT_SIZE_HEADER)
    _add_run_styled(p, tipo, bold=True, font_size=FONT_SIZE_HEADER)
    _set_cell_vertical_alignment(info_cell)

    # Row 2: Student info
    _merge_cells_in_row(table, 2, 0, 3)
    student_cell = table.rows[2].cells[0]
    p = student_cell.paragraphs[0]
    _add_run_styled(p, "Aluno(a): ________________________________", font_size=FONT_SIZE_HEADER)
    _add_run_styled(p, "     Turma: ______", font_size=FONT_SIZE_HEADER)
    _add_run_styled(p, "     Data: ____/____/________", font_size=FONT_SIZE_HEADER)
    _set_cell_vertical_alignment(student_cell)

    # Row 3: Professor, Peso, Nota
    _merge_cells_in_row(table, 3, 0, 3)
    prof_cell = table.rows[3].cells[0]
    p = prof_cell.paragraphs[0]

    profs = metadata.get("professors", [])
    prof_str = ", ".join(profs) if profs else "_______________"

    _add_run_styled(p, "Professor(a): ", font_size=FONT_SIZE_HEADER)
    _add_run_styled(p, prof_str, bold=True, font_size=FONT_SIZE_HEADER)
    _add_run_styled(p, "     Peso: ______", font_size=FONT_SIZE_HEADER)
    _add_run_styled(p, "     Nota: ______", font_size=FONT_SIZE_HEADER)
    _set_cell_vertical_alignment(prof_cell)

    # Set row heights
    for row in table.rows:
        tr = row._tr
        tr_pr = tr.get_or_add_trPr()
        tr_height = parse_xml(
            f'<w:trHeight {nsdecls("w")} w:val="400" w:hRule="atLeast"/>'
        )
        tr_pr.append(tr_height)

    return table


def _set_table_borders(table, sz="4", color="000000"):
    """Set uniform borders on entire table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>'
    )
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(borders)


# ---------------------------------------------------------------------------
# Formulário section
# ---------------------------------------------------------------------------

def _build_formulario(doc, formulario_frags: list, images: dict):
    """Add the formulário section (physics/math formulas)."""
    if not formulario_frags:
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, before=12, after=4)
    _add_run_styled(p, "FORMULÁRIO", bold=True, font_size=Pt(11))

    # Add a thin horizontal rule
    p_rule = doc.add_paragraph()
    p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p_rule, before=0, after=4)
    _add_run_styled(p_rule, "─" * 70, font_size=Pt(6), color="999999")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, before=2, after=2, line=14)

    for frag in formulario_frags:
        _render_fragment(p, frag, images, font_size=FONT_SIZE_SMALL)

    p_rule2 = doc.add_paragraph()
    p_rule2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p_rule2, before=4, after=8)
    _add_run_styled(p_rule2, "─" * 70, font_size=Pt(6), color="999999")


# ---------------------------------------------------------------------------
# Question rendering
# ---------------------------------------------------------------------------

def _render_fragment(paragraph, frag: dict, images: dict,
                     font_size=FONT_SIZE_BODY):
    """Render a single fragment (text, image, equation, line_break) into a paragraph."""
    ftype = frag.get("type", "text")

    if ftype == "text":
        _add_run_styled(
            paragraph,
            frag.get("content", ""),
            bold=frag.get("bold", False),
            italic=frag.get("italic", False),
            underline=frag.get("underline", False),
            font_size=font_size,
            color=frag.get("color"),
        )

    elif ftype == "equation":
        omml = frag.get("omml")
        if omml:
            _add_omml_to_paragraph(paragraph, omml)
        else:
            _add_run_styled(paragraph, frag.get("content", "[eq]"),
                          italic=True, font_size=font_size)

    elif ftype == "image":
        img_key = frag.get("image", "")
        img_b64 = images.get(img_key)
        if img_b64:
            img_bytes = base64.b64decode(img_b64)
            _add_image_to_paragraph(
                paragraph, img_bytes,
                width_px=frag.get("width_px"),
                height_px=frag.get("height_px"),
                crop=frag.get("crop"),
            )
        else:
            _add_run_styled(paragraph, f"[imagem: {img_key}]",
                          italic=True, font_size=font_size, color="FF0000")

    elif ftype == "line_break":
        paragraph.add_run().add_break()


def _render_fragments(paragraph, fragments: list, images: dict,
                      font_size=FONT_SIZE_BODY):
    """Render a list of fragments into a paragraph (inline, no image splitting)."""
    for frag in fragments:
        _render_fragment(paragraph, frag, images, font_size)


def _render_fragments_with_images(doc, paragraph, fragments: list, images: dict,
                                  font_size=FONT_SIZE_BODY):
    """Render fragments, breaking images into their own centered paragraphs.

    Text and equations go inline in the current paragraph; when an image is
    encountered, it gets its own centered paragraph, then a new paragraph
    continues the remaining text.  Returns the last paragraph used.
    """
    current_p = paragraph
    for frag in fragments:
        if frag.get("type") == "image":
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_spacing(p_img, before=4, after=4)
            _render_fragment(p_img, frag, images, font_size)
            current_p = doc.add_paragraph()
            _set_paragraph_spacing(current_p, before=2, after=2)
        else:
            _render_fragment(current_p, frag, images, font_size)
    return current_p


def _strip_number_prefix(fragments: list) -> list:
    """Remove the question number prefix (e.g. '3) ' or '3' + ') ') from fragments.

    Returns cleaned fragments without the leading number and parenthesis.
    Also moves any images that appear before the number to after it.
    """
    result = list(fragments)
    leading_images = []

    # Collect and remove leading non-text fragments (images before number)
    while result and result[0].get("type") != "text":
        if result[0].get("type") == "image":
            leading_images.append(result.pop(0))
        else:
            result.pop(0)

    if not result:
        return leading_images

    # Case 1: first text fragment is "N) text..." or "N)" alone
    first = result[0]
    content = first.get("content", "")
    m = re.match(r"^\s*\d{1,3}\s*\)\s*", content)
    if m:
        remaining = content[m.end():]
        if remaining.strip():
            result[0] = {**first, "content": remaining}
        else:
            result.pop(0)
        return leading_images + result

    # Case 2: first text is just "N" (number only), second starts with ")"
    m2 = re.match(r"^\s*\d{1,3}\s*$", content)
    if m2 and len(result) > 1:
        result.pop(0)
        # Remove the ") " from the next text fragment
        while result and result[0].get("type") != "text":
            if result[0].get("type") == "image":
                leading_images.append(result.pop(0))
            else:
                result.pop(0)
        if result:
            next_content = result[0].get("content", "")
            m3 = re.match(r"^\s*\)\s*", next_content)
            if m3:
                remaining = next_content[m3.end():]
                if remaining.strip():
                    result[0] = {**result[0], "content": remaining}
                else:
                    result.pop(0)
        return leading_images + result

    return leading_images + result


def _build_question(doc, question: dict, images: dict):
    """Render a single question with its alternatives."""
    stmt = question.get("statement", [])

    # Strip old number prefix and get clean content
    clean_frags = _strip_number_prefix(list(stmt))

    # Separate leading images (those that appeared before number in original)
    leading_images = []
    body_frags = []
    collecting_leading = True
    for f in clean_frags:
        if collecting_leading and f.get("type") == "image":
            leading_images.append(f)
        else:
            collecting_leading = False
            body_frags.append(f)

    # Render leading images in their own centered paragraph
    if leading_images:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(p_img, before=8, after=4)
        for img_frag in leading_images:
            _render_fragment(p_img, img_frag, images)

    # Question number + statement body
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=10 if not leading_images else 2, after=4)

    _add_run_styled(p, f"{question['number']}) ", bold=True,
                   font_size=FONT_SIZE_QUESTION_NUM)
    _render_fragments_with_images(doc, p, body_frags, images)

    # Source tag and weight as subtle annotation
    annotations = []
    if question.get("source_tag"):
        annotations.append(question["source_tag"])
    if question.get("weight"):
        annotations.append(f"(vale {question['weight']})")

    if annotations:
        p_annot = doc.add_paragraph()
        _set_paragraph_spacing(p_annot, before=0, after=2)
        _add_run_styled(p_annot, " ".join(annotations),
                       italic=True, font_size=FONT_SIZE_SMALL, color="666666")

    # Alternatives
    for alt in question.get("alternatives", []):
        p_alt = doc.add_paragraph()
        p_alt.paragraph_format.left_indent = Cm(1.0)
        _set_paragraph_spacing(p_alt, before=1, after=1)

        _add_run_styled(p_alt, f"{alt['letter']}) ", bold=True,
                       font_size=FONT_SIZE_BODY)
        _render_fragments(p_alt, alt.get("fragments", []), images)


# ---------------------------------------------------------------------------
# Gabarito grid
# ---------------------------------------------------------------------------

def _build_gabarito_grid(doc, questions: list):
    """
    Build the answer grid table at the end of the exam.
    Layout: rows of 5 questions each, with columns for each alternative.
    """
    if not questions:
        return

    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=16, after=6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run_styled(p, "GABARITO", bold=True, font_size=Pt(12))

    n_alts = 5  # a-e
    cols_per_q = n_alts + 1  # question number + alternatives
    questions_per_row = 5

    # Split questions into rows of 5
    q_groups = []
    for i in range(0, len(questions), questions_per_row):
        q_groups.append(questions[i:i + questions_per_row])

    for group in q_groups:
        # Header row + data row per group
        n_cols = len(group) * cols_per_q
        table = doc.add_table(rows=2, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(table, sz="4")

        # Header row: Q numbers spanning their columns
        for qi, q in enumerate(group):
            start_col = qi * cols_per_q
            # Question number cell
            cell = table.rows[0].cells[start_col]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run_styled(p, str(q["number"]), bold=True,
                          font_size=FONT_SIZE_SMALL)
            _set_cell_shading(cell, "D9E2F3")
            _set_cell_vertical_alignment(cell)

            # Alternative header cells
            for ai, letter in enumerate(ALT_LETTERS):
                col_idx = start_col + ai + 1
                if col_idx < n_cols:
                    cell = table.rows[0].cells[col_idx]
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _add_run_styled(p, letter.upper(), bold=True,
                                  font_size=FONT_SIZE_SMALL)
                    _set_cell_shading(cell, "D9E2F3")
                    _set_cell_vertical_alignment(cell)

        # Data row: empty cells for marking
        for qi, q in enumerate(group):
            start_col = qi * cols_per_q
            # Empty number cell
            cell = table.rows[1].cells[start_col]
            _set_cell_shading(cell, "F2F2F2")

            for ai in range(n_alts):
                col_idx = start_col + ai + 1
                if col_idx < n_cols:
                    cell = table.rows[1].cells[col_idx]
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _set_cell_vertical_alignment(cell)

        # Set row heights
        for row in table.rows:
            tr = row._tr
            tr_pr = tr.get_or_add_trPr()
            tr_height = parse_xml(
                f'<w:trHeight {nsdecls("w")} w:val="350" w:hRule="atLeast"/>'
            )
            tr_pr.append(tr_height)

        # Small space between groups
        p_space = doc.add_paragraph()
        _set_paragraph_spacing(p_space, before=2, after=2)


# ---------------------------------------------------------------------------
# Gabarito de respostas (answer key document)
# ---------------------------------------------------------------------------

def _build_answer_key(doc, questions: list, metadata: dict):
    """Build a separate answer key document (for professor use)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run_styled(p, "GABARITO — RESPOSTAS", bold=True, font_size=Pt(14))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subject = metadata.get("subject", "")
    series = metadata.get("series", "")
    _add_run_styled(p, f"{subject} — {series}", font_size=Pt(11))
    _set_paragraph_spacing(p, before=4, after=12)

    # Table: Questão | Resposta
    table = doc.add_table(rows=len(questions) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)

    # Header
    for ci, label in enumerate(["Questão", "Resposta"]):
        cell = table.rows[0].cells[ci]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run_styled(p, label, bold=True, font_size=FONT_SIZE_HEADER)
        _set_cell_shading(cell, "D9E2F3")
        _set_cell_vertical_alignment(cell)

    for ri, q in enumerate(questions):
        cell_num = table.rows[ri + 1].cells[0]
        p = cell_num.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run_styled(p, str(q["number"]), font_size=FONT_SIZE_HEADER)
        _set_cell_vertical_alignment(cell_num)

        cell_resp = table.rows[ri + 1].cells[1]
        p = cell_resp.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run_styled(p, "____", font_size=FONT_SIZE_HEADER)
        _set_cell_vertical_alignment(cell_resp)


# ---------------------------------------------------------------------------
# Main generator
# ---------------------------------------------------------------------------

def generate_exam(data: dict, output_path: str, logo_path: Optional[str] = None,
                  answer_key_path: Optional[str] = None):
    """
    Generate a standardized .docx exam from extraction JSON.

    Args:
        data: JSON from extractor.py (single or merged extraction)
        output_path: path for the output .docx
        logo_path: optional path to school logo image
        answer_key_path: optional path for separate answer key .docx
    """
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = PAGE_MARGIN_TOP
    section.bottom_margin = PAGE_MARGIN_BOTTOM
    section.left_margin = PAGE_MARGIN_LEFT
    section.right_margin = PAGE_MARGIN_RIGHT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_BODY
    font.size = FONT_SIZE_BODY

    # Get data
    metadata = data.get("metadata", {})
    formulario = data.get("formulario", [])
    questions = data.get("questions", [])
    images = data.get("images", {})

    # 1. Header table
    _build_header_table(doc, metadata, logo_path)

    # 2. Formulário (if present)
    if formulario:
        _build_formulario(doc, formulario, images)

    # 3. Questions
    for q in questions:
        _build_question(doc, q, images)

    # 4. Gabarito grid
    _build_gabarito_grid(doc, questions)

    # Save
    doc.save(output_path)
    print(f"Prova gerada: {output_path}")

    # 5. Separate answer key (if requested)
    if answer_key_path:
        ak_doc = Document()
        ak_section = ak_doc.sections[0]
        ak_section.top_margin = PAGE_MARGIN_TOP
        ak_section.bottom_margin = PAGE_MARGIN_BOTTOM
        ak_section.left_margin = PAGE_MARGIN_LEFT
        ak_section.right_margin = PAGE_MARGIN_RIGHT

        ak_style = ak_doc.styles["Normal"]
        ak_style.font.name = FONT_BODY
        ak_style.font.size = FONT_SIZE_BODY

        _build_answer_key(ak_doc, questions, metadata)
        ak_doc.save(answer_key_path)
        print(f"Gabarito gerado: {answer_key_path}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Gera prova .docx padronizada a partir do JSON do extrator"
    )
    parser.add_argument("input", help="Arquivo JSON de entrada (do extractor.py)")
    parser.add_argument("-o", "--output", required=True,
                       help="Arquivo .docx de saída")
    parser.add_argument("--logo", help="Imagem do logo do colégio (PNG/JPG)")
    parser.add_argument("--gabarito",
                       help="Gera gabarito separado no caminho especificado")

    args = parser.parse_args()

    with open(args.input, "r", encoding="utf-8") as f:
        data = json.load(f)

    generate_exam(data, args.output, logo_path=args.logo,
                  answer_key_path=args.gabarito)


if __name__ == "__main__":
    main()
